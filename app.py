"""
AI 토론
- Streamlit(sync) UI + Playwright(async) 자동화 + pyngrok 터널링
- 실행: streamlit run app.py
"""

from __future__ import annotations

import asyncio
import ctypes
import json
import os
import re
import threading
import time
import urllib.parse
import urllib.request
from datetime import datetime
from io import BytesIO
from pathlib import Path
from queue import Empty, Queue
from typing import Optional

import streamlit as st
from dotenv import load_dotenv
from playwright.async_api import (
    Page,
    TimeoutError as PWTimeout,
    async_playwright,
)
from pyngrok import ngrok

# Word 문서 생성용
from docx import Document
from docx.shared import Pt, RGBColor

load_dotenv()

# ──────────────────────────── 설정 ────────────────────────────
CHROME_USER_DATA_DIR = (
    os.getenv("CHROME_USER_DATA_DIR", "").strip()
    or str(Path.home() / "ai_debate_profile")
)
NGROK_AUTHTOKEN = os.getenv("NGROK_AUTHTOKEN", "").strip()
STREAMLIT_PORT = int(os.getenv("STREAMLIT_PORT", "8501"))
BROWSER_CHANNEL  = os.getenv("BROWSER_CHANNEL", "chrome").strip() or None
GOOGLE_ACCOUNT   = os.getenv("GOOGLE_ACCOUNT", "").strip()

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
TELEGRAM_CHAT_ID = os.getenv("TELEGRAM_CHAT_ID", "").strip()

def set_windows_clipboard_text(text: str) -> None:
    """Set Unicode text on the Windows clipboard without extra dependencies."""
    if os.name != "nt":
        raise RuntimeError("Windows clipboard API is unavailable on this platform")
    user32 = ctypes.windll.user32
    kernel32 = ctypes.windll.kernel32
    user32.OpenClipboard.argtypes = [ctypes.c_void_p]
    user32.OpenClipboard.restype = ctypes.c_int
    user32.EmptyClipboard.restype = ctypes.c_int
    user32.SetClipboardData.argtypes = [ctypes.c_uint, ctypes.c_void_p]
    user32.SetClipboardData.restype = ctypes.c_void_p
    user32.CloseClipboard.restype = ctypes.c_int
    kernel32.GlobalAlloc.argtypes = [ctypes.c_uint, ctypes.c_size_t]
    kernel32.GlobalAlloc.restype = ctypes.c_void_p
    kernel32.GlobalLock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalLock.restype = ctypes.c_void_p
    kernel32.GlobalUnlock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalUnlock.restype = ctypes.c_int
    kernel32.GlobalFree.argtypes = [ctypes.c_void_p]
    kernel32.GlobalFree.restype = ctypes.c_void_p

    data = (text + "\0").encode("utf-16le")
    handle = kernel32.GlobalAlloc(0x0002, len(data))  # GMEM_MOVEABLE
    if not handle:
        raise RuntimeError("Windows clipboard memory allocation failed")

    locked = kernel32.GlobalLock(handle)
    if not locked:
        kernel32.GlobalFree(handle)
        raise RuntimeError("Windows clipboard memory lock failed")

    ctypes.memmove(locked, data, len(data))
    kernel32.GlobalUnlock(handle)

    opened = False
    for _ in range(10):
        if user32.OpenClipboard(None):
            opened = True
            break
        time.sleep(0.05)
    if not opened:
        kernel32.GlobalFree(handle)
        raise RuntimeError("Windows clipboard is busy")

    try:
        user32.EmptyClipboard()
        if not user32.SetClipboardData(13, handle):  # CF_UNICODETEXT
            kernel32.GlobalFree(handle)
            raise RuntimeError("Windows clipboard SetClipboardData failed")
        handle = None
    finally:
        user32.CloseClipboard()


def get_windows_clipboard_text() -> str:
    """Read Unicode text from the Windows clipboard without extra dependencies."""
    user32 = ctypes.windll.user32
    kernel32 = ctypes.windll.kernel32
    user32.OpenClipboard.argtypes = [ctypes.c_void_p]
    user32.OpenClipboard.restype = ctypes.c_int
    user32.IsClipboardFormatAvailable.argtypes = [ctypes.c_uint]
    user32.IsClipboardFormatAvailable.restype = ctypes.c_int
    user32.GetClipboardData.argtypes = [ctypes.c_uint]
    user32.GetClipboardData.restype = ctypes.c_void_p
    user32.CloseClipboard.restype = ctypes.c_int
    kernel32.GlobalLock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalLock.restype = ctypes.c_void_p
    kernel32.GlobalUnlock.argtypes = [ctypes.c_void_p]
    kernel32.GlobalUnlock.restype = ctypes.c_int

    opened = False
    for _ in range(10):
        if user32.OpenClipboard(None):
            opened = True
            break
        time.sleep(0.05)
    if not opened:
        return ""

    try:
        if not user32.IsClipboardFormatAvailable(13):  # CF_UNICODETEXT
            return ""
        handle = user32.GetClipboardData(13)
        if not handle:
            return ""
        locked = kernel32.GlobalLock(handle)
        if not locked:
            return ""
        try:
            return ctypes.wstring_at(locked)
        finally:
            kernel32.GlobalUnlock(handle)
    finally:
        user32.CloseClipboard()


AI_URLS = {
    "perplexity": "https://www.perplexity.ai/",
    "claude":  "https://claude.ai/new",
    "chatgpt": "https://chatgpt.com/",
    "gemini":  "https://gemini.google.com/app",
}

DEFAULT_DEBATE_ORDER = ["gemini", "chatgpt", "claude"]
DEBATE_ORDER_OPTIONS = {
    "Gemini → ChatGPT → Claude": ["gemini", "chatgpt", "claude"],
}
# ── AutoGen 패턴 1: 에이전트 페르소나 (AssistantAgent의 system_message 역할) ──
AGENT_PERSONAS = {
    "perplexity": (
        "너는 토론자가 아니라 검색 비서이자 근거 수집 담당이야. "
        "판단과 추천을 최소화하고 최신 사실, 공식 출처, 확인일, 원문 링크를 수집해."
    ),
    "gemini": (
        "너는 교차검증자이자 근거 기반 발산 탐색자야. Perplexity 팩트팩의 핵심 사실을 "
        "검증됨·상충함·미검증으로 분류하고, 가능하면 독립된 출처로 교차검증해. "
        "그 다음 누락된 사실, 반대 근거, 숨은 변수와 대안을 확장해. "
        "새로운 사실이나 수치를 제시할 때는 원문 URL을 붙이고, 출처가 없으면 사실로 "
        "단정하지 말고 반드시 '가설' 또는 '추가 확인 필요'로 표시해. "
        "최종 결론은 내리지 말고 사실 검증과 가능성 확장에 집중해."
    ),
    "chatgpt": (
        "너는 논리와 현실성 검증자야. Gemini 또는 이전 화자의 주장에 있는 논리적 허점, "
        "근거 부족, 비용·시간·운영 리스크, 실행 가능성 과대평가를 집중적으로 검토해. "
        "Perplexity와 Gemini가 같은 원출처를 반복 인용한 것은 독립 교차검증으로 인정하지 마. "
        "링크가 실제 주장과 일치하는지, 출처가 공식·독립·제휴성 자료 중 무엇인지 구분해. "
        "좋아 보이는 결론일수록 현실에서 깨질 지점을 먼저 찾아."
    ),
    "claude": (
        "너는 수렴 판단자야. Gemini의 발산과 GPT의 검증을 종합해 최종 판단과 실행안을 제시해. "
        "무조건 중립을 피하고, 어떤 선택이 더 타당한지 명확히 판단해. "
        "다만 GPT가 지적한 리스크를 무시하지 말고 결론의 한계와 조건을 함께 밝혀."
    ),
}

# ── AutoGen 패턴 3: 종료 조건 키워드 (TextMentionTermination 역할) ──
CONSENSUS_KEYWORDS = [
    "모두 동의", "완전히 동의", "합의에 이르", "더 이상 반박할",
    "의견 일치", "논쟁이 불필요", "충분한 합의",
]

# 자주 바뀌는 클래스 대신 placeholder/aria-label/data-testid 기반 셀렉터
SELECTORS = {
    "perplexity": {
        "input": (
            'textarea[placeholder*="Ask" i], textarea[aria-label*="Ask" i], '
            '[contenteditable="true"][data-lexical-editor="true"], '
            'div[contenteditable="true"], textarea'
        ),
        "send": (
            'button[aria-label*="Submit" i], button[aria-label*="Send" i], '
            'button[type="submit"], button[data-testid*="submit" i]'
        ),
        "stop": (
            'button[aria-label*="Stop" i], button[data-testid*="stop" i], '
            'button:has-text("Stop")'
        ),
        "response": (
            '[data-testid*="answer" i], main [class*="prose"], '
            'main [class*="markdown"], article'
        ),
    },
    "claude": {
        "input":    'div[contenteditable="true"][aria-label], div.ProseMirror[contenteditable="true"]',
        "send":     'button[aria-label="Send Message"], button[aria-label="Send message"], button[aria-label*="Send" i]',
        "stop":     'button[aria-label="Stop Response"], button[aria-label*="Stop" i]',
        # 최신 Claude UI: 응답 본문 컨테이너는 div.standard-markdown.
        # 단락 클래스 font-claude-response-body 의 부모로도 잡힘.
        "response": 'div.standard-markdown, [data-testid="assistant-message"], div.font-claude-message, .prose',
    },
    "chatgpt": {
        "input":    '#prompt-textarea, div[contenteditable="true"]#prompt-textarea',
        "send":     'button[data-testid="send-button"], button[aria-label*="Send" i]',
        "stop":     'button[data-testid="stop-button"], button[aria-label*="Stop" i]',
        "response": 'div[data-message-author-role="assistant"]',
    },
    "gemini": {
        "input":    'rich-textarea div[contenteditable="true"], div.ql-editor[contenteditable="true"]',
        "send":     'button[aria-label="Send message"], button[aria-label*="Send" i], button[aria-label*="Submit" i], button.send-button, button:has(mat-icon:has-text("send"))',
        "stop":     'button[aria-label="Stop response"], button.stop, button[aria-label*="Stop" i]',
        "response": 'message-content, model-response, .model-response-text',
    },
}

# ──────────────────────────── 작업자 ────────────────────────────
class DebateWorker:
    """별도 스레드에서 자체 asyncio 루프를 굴리며 Playwright를 운용."""

    def __init__(self) -> None:
        self.cmd_queue: Queue = Queue()
        self.inject_queue: Queue = Queue()   # Human-in-the-loop 개입 메시지
        self.messages: list[dict] = []
        self.current_debate_started_at: float = 0.0
        self.status: str = "booting"
        self.error: Optional[str] = None
        self.lock = threading.Lock()
        self.thread: Optional[threading.Thread] = None
        self._loop: Optional[asyncio.AbstractEventLoop] = None
        self.pages: dict[str, Page] = {}
        self.last_docx: Optional[dict] = None
        self.persist_records: bool = False

    def inject_message(self, msg: str) -> None:
        """진행 중인 토론에 사용자 개입 메시지를 주입."""
        self.inject_queue.put(msg)

    # ── 외부에서 호출하는 동기 API ──
    def start(self) -> None:
        if self.thread and self.thread.is_alive():
            return
        self.thread = threading.Thread(target=self._thread_main, daemon=True)
        self.thread.start()

    def submit_topic(
        self,
        topic: str,
        rounds: int,
        debate_order: Optional[list[str]] = None,
        no_local_records: bool = True,
    ) -> None:
        # 메인 스레드에서 즉시 running으로 (워커 스레드 픽업까지 기다리지 않게)
        with self.lock:
            if self.status == "ready":
                self.status = "running"
                self.messages = []
                self.current_debate_started_at = time.time()
                self.last_docx = None
                self.error = None
                self.persist_records = not no_local_records
        self._clear_inject_queue()
        self.cmd_queue.put({
            "type": "debate",
            "topic": topic,
            "rounds": rounds,
            "debate_order": normalize_debate_order(debate_order),
            "no_local_records": no_local_records,
        })

    def submit_summary(self, summarizer: str = "gemini") -> None:
        self.cmd_queue.put({"type": "summarize", "summarizer": summarizer})

    def snapshot(self) -> tuple[str, Optional[str], list[dict], Optional[dict]]:
        with self.lock:
            return self.status, self.error, list(self.messages), self.last_docx

    # ── 내부 ──
    def _post(self, role: str, content: str) -> None:
        with self.lock:
            self.messages.append({"role": role, "content": content, "ts": time.time()})
            persist_records = self.persist_records
        if not persist_records:
            return
        try:
            log_dir = Path(__file__).parent / "logs"
            log_dir.mkdir(exist_ok=True)
            ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            preview = content.replace("\r", " ").replace("\n", " ")[:1200]
            with (log_dir / "debate_events.log").open("a", encoding="utf-8") as f:
                f.write(f"[{ts}] {role}: {preview}\n")
        except Exception:
            pass

    def _set_status(self, s: str) -> None:
        with self.lock:
            self.status = s

    def _clear_inject_queue(self) -> None:
        try:
            while True:
                self.inject_queue.get_nowait()
        except Empty:
            pass
        except Exception:
            pass

    def _current_debate_messages(self, roles: Optional[set[str]] = None) -> list[dict]:
        with self.lock:
            started_at = self.current_debate_started_at
            messages = list(self.messages)
        if started_at:
            messages = [m for m in messages if float(m.get("ts", 0)) >= started_at]
        if roles:
            messages = [m for m in messages if m.get("role") in roles]
        return messages

    def _fail(self, msg: str) -> None:
        with self.lock:
            self.status = "error"
            self.error = msg
        self._post("system", f"❌ {msg}")

    def _thread_main(self) -> None:
        self._loop = asyncio.new_event_loop()
        asyncio.set_event_loop(self._loop)
        try:
            self._loop.run_until_complete(self._async_main())
        except Exception as e:  # 최후의 안전망
            self._fail(f"작업자 스레드 비정상 종료: {e}")

    async def _async_main(self) -> None:
        async with async_playwright() as pw:
            try:
                Path(CHROME_USER_DATA_DIR).mkdir(parents=True, exist_ok=True)
                ctx = await pw.chromium.launch_persistent_context(
                    user_data_dir=CHROME_USER_DATA_DIR,
                    channel=BROWSER_CHANNEL,
                    headless=False,
                    viewport={"width": 1320, "height": 920},
                    args=[
                        "--disable-blink-features=AutomationControlled",
                        "--start-maximized",
                    ],
                )
            except Exception as e:
                self._fail(
                    "크롬 실행 실패: "
                    f"{e}. Chrome이 설치되어 있는지 확인하거나, `.env`에서 "
                    "`BROWSER_CHANNEL=`처럼 비워 둔 뒤 `python -m playwright install chromium`을 실행하세요. "
                    "같은 프로필 폴더를 다른 크롬 창이 사용 중이어도 실행이 실패할 수 있습니다."
                )
                return

            try:
                for name, url in AI_URLS.items():
                    page = await ctx.new_page()
                    await page.goto(url, wait_until="domcontentloaded", timeout=60_000)
                    self.pages[name] = page
                self._post(
                    "system",
                    "✅ Perplexity / Gemini / ChatGPT / Claude 탭을 열었습니다. "
                    "각 사이트에 **로그인** 되어 있는지 크롬 창에서 확인하세요.",
                )
            except Exception as e:
                self._fail(f"AI 사이트 탭 오픈 실패: {e}")
                return

            self._set_status("ready")

            while True:
                try:
                    cmd = self.cmd_queue.get_nowait()
                except Empty:
                    await asyncio.sleep(0.4)
                    continue

                if cmd["type"] == "debate":
                    try:
                        await self._run_debate(
                            cmd["topic"],
                            cmd["rounds"],
                            normalize_debate_order(cmd.get("debate_order")),
                        )
                    except Exception as e:
                        self._fail(f"토론 도중 오류: {e}")
                        self._set_status("ready")  # 다음 주제는 받을 수 있게

    # ── 토론 로직 ──
    async def _run_debate(self, topic: str, rounds: int, debate_order: list[str]) -> None:
        """
        0단계: Perplexity가 최신 정보와 출처를 팩트팩으로 수집
        모든 세트: Gemini → ChatGPT → Claude
        마지막 세트의 마지막 AI가 최종 종합 글 작성 → GPT가 Word 파일 생성 → Telegram 알림

        rounds=1 → 선택 순서 × 1 세트 (마지막 AI = 종합)
        rounds=2 → 선택 순서 × 2 세트 (마지막 AI = 종합)
        """
        self._set_status("running")
        with self.lock:
            if not self.current_debate_started_at:
                self.current_debate_started_at = time.time()
            self.messages = []
            self.last_docx = None
        self._post("user", f"📌 **주제**\n\n{topic}")
        self._post(
            "system",
            "🧭 진행 순서: PERPLEXITY 근거 수집 → "
            f"{' → '.join(ai.upper() for ai in debate_order)} → "
            "GPT 최종 반대자 검토 → CLAUDE 수정 최종안",
        )

        total_sets = rounds
        last_speaker: Optional[str] = None
        last_reply: str = topic
        final_text = ""
        fact_pack = ""
        report_title = make_report_title(topic)
        report_filename = make_report_filename(topic)

        research_prompt = (
            f"{AGENT_PERSONAS['perplexity']}\n\n"
            "아래 주제에 대해 최신 정보를 검색해 팩트팩을 만들어줘.\n"
            "1. 공식 출처를 최우선으로 사용해.\n"
            "2. 제조사·기관·공식 도움말·보도자료·논문·공식 판매 페이지를 먼저 확인해.\n"
            "3. 블로그·커뮤니티·사용자 리뷰는 보조 근거로 분리해.\n"
            "4. 핵심 사실을 표 또는 항목으로 정리해.\n"
            "5. 각 사실마다 출처 이름, 원문 URL, 확인일을 붙여.\n"
            "6. 출처에서 직접 확인되지 않은 숫자나 주장은 만들지 마.\n"
            "7. 결론이나 추천은 최소화하고 팩트 수집에 집중해.\n\n"
            f"주제:\n{topic}"
        )
        self._post("system", "🔎 Perplexity가 최신 정보와 공식 출처를 수집하는 중…")
        try:
            fact_pack = await self._ask("perplexity", research_prompt)
            self._post("perplexity", "🔎 **Perplexity 팩트팩**\n\n" + fact_pack)
        except Exception as e:
            self._post(
                "system",
                "⚠️ Perplexity 근거 수집에 실패했지만 토론은 계속 진행합니다. "
                f"Chrome 창에서 보안 확인/로그인이 필요한지 확인하세요: {e}",
            )
            fact_pack = (
                "Perplexity 팩트팩을 확보하지 못했습니다. 모든 모델은 최신 사실과 수치를 "
                "확정적으로 단정하지 말고 추가 확인 필요 여부를 명시해야 합니다."
            )

        fact_pack_context = (
            "\n\n=== Perplexity 팩트팩 ===\n"
            f"{fact_pack}\n"
            "=== 팩트팩 사용 규칙 ===\n"
            "- 자료를 그대로 믿지 말고 출처 신뢰도와 문장-출처 일치 여부를 검토할 것.\n"
            "- 공식 출처와 리뷰·커뮤니티 의견을 분리할 것.\n"
            "- 출처 없는 수치나 사실은 확정적으로 사용하지 말 것.\n"
        )

        gemini_verification_instruction = (
            "Perplexity 팩트팩을 교차검증하고 부족한 사실을 추가 탐색해.\n"
            "다음 형식을 반드시 지켜:\n"
            "1. 핵심 사실 검증표: 주장 / 상태(검증됨·상충함·미검증) / 근거 / URL\n"
            "2. Perplexity와 상충하는 자료\n"
            "3. 누락된 추가 팩트와 원문 URL\n"
            "4. 사실과 분리된 가설·대안·숨은 변수\n"
            "5. 추가 확인이 필요한 항목\n"
            "같은 원출처를 재인용한 것은 독립 교차검증이라고 부르지 마. "
            "출처 URL이 없는 새 수치나 사실은 만들지 말고 최종 결론도 내리지 마.\n"
        )

        for s in range(total_sets):
            is_final_set = (s == total_sets - 1)
            label = f"최종 세트 ({s+1}/{total_sets})" if is_final_set else f"세트 {s+1}/{total_sets}"
            self._post("system", f"🔁 {label} 시작")

            for ai in debate_order:
                human_inject = self._pop_inject()
                persona = f"[너의 역할]: {AGENT_PERSONAS.get(ai, '')}\n\n"
                is_final_speaker = is_final_set and ai == debate_order[-1]

                if last_speaker is None:
                    prompt = (
                        f"{persona}"
                        f"주제: {topic}\n\n"
                        f"{gemini_verification_instruction}"
                        f"{fact_pack_context}"
                        f"{human_inject}"
                    )
                elif is_final_speaker:
                    # 마지막 AI: 전체 토론 종합 글 작성 (Word 변환은 GPT가 담당)
                    convo = self._current_debate_messages(
                        {"perplexity", "chatgpt", "gemini", "claude"}
                    )
                    transcript = "\n\n".join(
                        f"[{m['role'].upper()}]\n{m['content']}" for m in convo
                    )
                    prompt = (
                        f"{persona}"
                        f"아래는 '{topic}'에 대한 전체 토론 기록이다.\n\n"
                        f"이 토론을 바탕으로 **{report_title}** 제목의 최종 종합 보고서를 한국어로 작성해줘.\n"
                        "- 마크다운 헤더(##/###), 표, 글머리표 적극 활용\n"
                        "- 핵심 결론 요약 (3~5문장)\n"
                        "- Perplexity 핵심 근거와 출처 신뢰도 요약\n"
                        "- Gemini 교차검증 결과의 검증됨·상충함·미검증 항목 요약\n"
                        "- GPT / Gemini / Claude 주장 비교표\n"
                        "- 통합 결론 + 실행 가능한 권고사항\n"
                        "- 확인된 사실, 추정, 추가 확인 필요 항목을 구분\n"
                        "- 최종 결론은 검증됨 항목을 중심으로 작성\n"
                        "- 서론 없이 바로 본론\n\n"
                        f"=== 토론 기록 ===\n{transcript}"
                        f"{human_inject}"
                    )
                else:
                    role_instruction = (
                        f"{gemini_verification_instruction}"
                        "앞선 토론에서 새로 나온 핵심 주장도 함께 재검토해."
                        if ai == "gemini"
                        else
                        "Perplexity와 Gemini 출처의 신뢰도, 독립성, 주장-원문 일치 여부를 확인하고 "
                        "상대 주장의 논리적 허점과 현실적 리스크를 검증해. "
                        "Gemini가 추가한 팩트도 검증 전에는 확정 사실로 취급하지 마."
                    )
                    prompt = (
                        f"{persona}"
                        f"주제: {topic}\n\n"
                        f"[{last_speaker.upper()}의 주장]\n{last_reply}\n\n"
                        f"{role_instruction} 한국어로 답해."
                        f"{fact_pack_context}"
                        f"{human_inject}"
                    )

                self._post("system", f"➡️ {ai.upper()} 에게 전달 중…")
                try:
                    reply = await self._ask(ai, prompt)
                except Exception as e:
                    if ai != "gemini":
                        raise
                    reply = (
                        "Gemini 웹 UI 자동화가 이번 세트에서 실패했습니다. "
                        "이후 최종 정리 단계에서 데이터 근거, 리스크, 실행 가능성 관점의 검토를 보완해야 합니다. "
                        f"실패 원인: {e}"
                    )
                    self._post(
                        "system",
                        "⚠️ GEMINI 자동화가 실패했지만 토론은 계속 진행합니다. "
                        "최종 정리 단계에서 Gemini 관점의 데이터·리스크 검토까지 보완하게 합니다.",
                    )

                if is_final_speaker:
                    self._post(ai, f"📝 **{ai.upper()} 최종 종합**\n\n" + reply)
                    final_text = reply
                else:
                    self._post(ai, reply)

                last_speaker = ai
                last_reply = reply

                # 종료 조건 (최종 세트 아닐 때만)
                if not is_final_set and any(kw in reply for kw in CONSENSUS_KEYWORDS):
                    self._post("system", "🤝 합의 신호 — 최종 세트로 건너뜁니다.")
                    # 현재 for ai 루프 종료 후 바깥 for s 루프에서 is_final_set=True 되도록
                    s = total_sets - 2  # type: ignore[assignment]
                    break

        # ── Claude 초안에 대한 GPT 최종 반대자 검토 → Claude 수정 최종안 ──
        draft_final_text = final_text
        if draft_final_text.strip():
            review_transcript = "\n\n".join(
                f"[{m['role'].upper()}]\n{m['content']}"
                for m in self._current_debate_messages(
                    {"perplexity", "chatgpt", "gemini", "claude"}
                )
            )
            final_review_prompt = (
                "당신은 최종 반대자 역할이다. 아래 Claude의 최종 결론이 틀렸을 가능성을 검토하라.\n\n"
                "특히 다음을 반드시 점검하라.\n"
                "1. 결론이 너무 성급한가?\n"
                "2. 중요한 리스크를 무시했는가?\n"
                "3. 실행 가능성이 과대평가되었는가?\n"
                "4. 반대 결론이 더 타당할 가능성이 있는가?\n"
                "5. 최종 결론을 수정해야 한다면 어떻게 수정해야 하는가?\n\n"
                "6. Gemini가 상충함·미검증으로 표시한 사실이 결론에서 확정 사실처럼 사용되었는가?\n"
                "7. Perplexity와 Gemini가 같은 원출처를 반복한 것을 독립 검증으로 잘못 계산했는가?\n\n"
                "비판을 위한 비판은 피하고, 실제 의사결정 품질을 높이는 반론만 제시하라.\n"
                "한국어로, 항목별로 명확하게 작성하라.\n\n"
                f"=== 주제 ===\n{topic}\n\n"
                f"=== 전체 토론 기록 ===\n{review_transcript}\n\n"
                f"=== Claude 최종 결론 초안 ===\n{draft_final_text}"
            )
            self._post("system", "🧪 GPT에게 최종 반대자 검토 요청 중…")
            try:
                final_review = await self._ask("chatgpt", final_review_prompt)
                self._post("chatgpt", "🧪 **GPT 최종 반대자 검토**\n\n" + final_review)

                revision_prompt = (
                    "너는 최종 수렴 판단자다. 아래 Claude 초안과 GPT 최종 반대자 검토를 바탕으로 "
                    "최종 보고서를 수정하라.\n\n"
                    "규칙:\n"
                    "- GPT의 지적 중 타당한 것은 결론과 실행안에 반영하라.\n"
                    "- 반영하지 않는 지적은 왜 반영하지 않는지 짧게 밝혀라.\n"
                    "- 최종 결론을 더 보수적으로 수정해야 하면 수정하라.\n"
                    "- 실행 조건, 리스크, 반대 결론 가능성을 명확히 남겨라.\n"
                    "- 상충함·미검증 상태의 사실은 결론 근거에서 제외하거나 불확실성을 명시하라.\n"
                    "- 서론 없이 바로 본론으로 작성하라.\n\n"
                    f"=== 주제 ===\n{topic}\n\n"
                    f"=== Claude 초안 ===\n{draft_final_text}\n\n"
                    f"=== GPT 최종 반대자 검토 ===\n{final_review}"
                )
                self._post("system", "🧭 Claude에게 GPT 검토 반영 수정 최종안 요청 중…")
                revised_final = await self._ask("claude", revision_prompt)
                final_text = revised_final
                self._post("claude", "📝 **CLAUDE 수정 최종안**\n\n" + revised_final)
            except Exception as e:
                self._post(
                    "system",
                    f"⚠️ 최종 반대자 검토/수정 단계 실패 — Claude 초안을 최종안으로 사용합니다: {e}",
                )
                final_text = draft_final_text

        # ── GPT에게 Word 파일 생성 요청 ──
        self._post("system", "📄 GPT에게 Word 파일 생성 요청 중…")
        word_prompt = (
            "아래 보고서를 실제 .docx 파일로 만들어줘.\n"
            f"문서 제목은 반드시 `{report_title}`로 하고, "
            f"파일명은 `{report_filename}`로 만들어줘.\n"
            "Code Interpreter를 반드시 사용해서 python-docx로 "
            "헤더·표·글머리표가 적용된 Word 파일을 생성하고, "
            "다운로드 가능하게 첨부해줘.\n\n"
            f"=== 보고서 내용 ===\n{final_text}"
        )
        await self._ask("chatgpt", word_prompt)
        self._post("system", "⏳ GPT Word 파일 다운로드 시도 중…")

        # GPT 첨부 파일 + 링크 캡처 시도
        gpt_docx, gpt_link = await self._try_download_chatgpt_attachment()

        # ChatGPT 대화 링크 (문서를 만든 대화 자체의 URL)
        gpt_chat_url: Optional[str] = None
        try:
            url = self.pages["chatgpt"].url
            if url and ("chatgpt.com/c/" in url or "chatgpt.com/share/" in url):
                gpt_chat_url = url
        except Exception:
            pass

        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        out_dir = Path(__file__).parent / "out"
        with self.lock:
            persist_records = self.persist_records
        if persist_records:
            out_dir.mkdir(exist_ok=True)

        if gpt_docx:
            docx_bytes = gpt_docx
            filename = f"{Path(report_filename).stem}_gpt_{stamp}.docx"
            if persist_records:
                (out_dir / filename).write_bytes(docx_bytes)
            with self.lock:
                self.last_docx = {
                    "bytes": docx_bytes, "filename": filename, "topic": topic,
                    "title": report_title, "gpt_link": gpt_link, "gpt_chat_url": gpt_chat_url,
                    "persisted": persist_records,
                }
            if persist_records:
                self._post("system", f"✅ GPT Word 파일 저장 → `out/{filename}`")
            else:
                self._post("system", "✅ GPT Word 파일 생성 완료 — 로컬 디스크에는 저장하지 않습니다.")
            if gpt_link:
                self._post("system", f"🔗 ChatGPT 생성 문서 링크: {gpt_link}")
            if gpt_chat_url:
                self._post("system", f"💬 문서를 만든 ChatGPT 대화: {gpt_chat_url}")
        else:
            # 폴백: 마크다운 → Word 변환 (Python)
            self._post("system", "ℹ️ GPT 파일 캡처 실패 — 마크다운→Word 변환 사용")
            try:
                docx_bytes = build_docx_from_markdown(topic, final_text, report_title=report_title)
                filename = f"{Path(report_filename).stem}_{stamp}.docx"
                if persist_records:
                    (out_dir / filename).write_bytes(docx_bytes)
                with self.lock:
                    self.last_docx = {
                        "bytes": docx_bytes, "filename": filename, "topic": topic,
                        "title": report_title, "gpt_link": gpt_link, "gpt_chat_url": gpt_chat_url,
                        "persisted": persist_records,
                    }
            except Exception as e:
                self._post("system", f"⚠️ Word 생성 실패: {e}")
                docx_bytes = None
                filename = None

        # ── 완료 알림 발송: 본문/파일은 싣지 않고 완료 사실만 알림 ──
        try:
            sent_channels = send_completion_notifications(
                topic,
                filename,
                gpt_chat_url,
                persist_records,
            )
            if sent_channels:
                self._post("system", f"📣 완료 알림 발송 → {', '.join(sent_channels)}")
            else:
                self._post("system", "ℹ️ 알림 채널 미설정 — 화면에서 완료 상태만 표시합니다.")
        except Exception as e:
            self._post("system", f"⚠️ 완료 알림 발송 실패: {e}")

        # ── 히스토리 저장 ──
        if persist_records:
            try:
                snap_msgs = self._current_debate_messages()
                saved = save_history(topic, snap_msgs, filename)
                self._post("system", f"💾 히스토리 저장됨 → `history/{saved}`")
            except Exception as e:
                self._post("system", f"⚠️ 히스토리 저장 실패: {e}")
        else:
            self._post(
                "system",
                "🔒 로컬 기록 없음 — 질문·답변, 실행 로그, Word 파일을 디스크에 저장하지 않았습니다.",
            )

        if persist_records:
            self._post("system", "🏁 **토론 완료** — Word 다운로드와 히스토리 확인이 가능합니다.")
        else:
            self._post(
                "system",
                "🏁 **토론 완료** — Word 다운로드가 가능합니다. "
                "서버 재시작 전까지 메모리에만 유지됩니다.",
            )
        self._set_status("ready")

    async def _summarize(self, summarizer: str) -> None:
        """전체 토론 내용을 지정 AI에게 보내 통합 요약 받기."""
        convo = self._current_debate_messages(
            {"user", "perplexity", "claude", "chatgpt", "gemini"}
        )
        if not convo:
            self._post("system", "⚠️ 요약할 토론 기록이 없습니다.")
            return

        transcript = "\n\n".join(f"=== {m['role'].upper()} ===\n{m['content']}" for m in convo)
        prompt = (
            "다음은 Perplexity 팩트팩과 Claude, ChatGPT, Gemini의 토론 기록이다.\n"
            "Perplexity는 의견이 아니라 근거 수집 결과로 취급하고, "
            "각 AI 주장의 공통점·차이점·핵심 갈등 지점을 비교 정리하고, "
            "마지막에 통합 결론을 제시해줘. "
            "한국어로, 마크다운 헤더/표/리스트를 적극 활용해 깔끔히 정리해줘.\n\n"
            f"{transcript}"
        )

        self._set_status("running")
        self._post("system", f"📝 {summarizer.upper()}에게 통합 요약 요청 중…")
        try:
            summary = await self._ask(summarizer, prompt)
        except Exception as e:
            self._set_status("ready")
            raise e
        self._post("summary", summary)
        self._post("system", "✅ 요약 완료 — 사이드바의 'Word 다운로드' 버튼으로 받을 수 있어요.")
        self._set_status("ready")

    async def _ask(self, ai: str, prompt: str) -> str:
        page = self.pages[ai]
        sel = SELECTORS[ai]
        await page.bring_to_front()

        # 0) 매번 "새 채팅"으로 이동 (이전 대화 컨텍스트 제거)
        try:
            await page.goto(AI_URLS[ai], wait_until="domcontentloaded", timeout=60_000)
        except Exception:
            pass  # 이미 그 페이지일 수도 있으니 실패해도 계속

        # 1) 입력창 확보. 없으면 로그인 화면일 가능성 → Google 로그인 자동 시도
        try:
            await self._dismiss_blocking_overlays(page)
            input_el = await page.wait_for_selector(sel["input"], timeout=8_000, state="visible")
        except PWTimeout:
            # "Continue with Google" 자동 클릭 시도 (best-effort)
            triggered = await self._try_google_signin(page, ai)
            if triggered:
                self._post(
                    "system",
                    f"🔐 {ai.upper()}: 'Continue with Google' 클릭함. "
                    f"Google 계정 선택/확인이 필요하면 크롬 창에서 진행해주세요."
                )
            else:
                self._post(
                    "system",
                    f"⚠️ {ai.upper()} 입력창을 찾지 못했습니다. **크롬 창에서 로그인/캡챠를 직접 해결**해 주세요. "
                    f"입력창이 보이면 자동으로 진행됩니다 (최대 5분 대기).",
                )
            # 사용자가 로그인 완료할 때까지 최대 5분 대기
            await self._dismiss_blocking_overlays(page)
            input_el = await page.wait_for_selector(sel["input"], timeout=300_000, state="visible")

        # 2) 기존 내용 제거 후 프롬프트 전체를 클립보드 붙여넣기로 삽입
        await self._safe_click(page, input_el, f"{ai.upper()} 입력창")
        await self._fill_prompt(page, input_el, prompt, ai)

        await asyncio.sleep(0.3)

        before_response = await self._last_response_text(page, sel, ai)

        # 3) 전송: 버튼 클릭 우선, 실패 시 Enter
        sent = False
        try:
            btn = None
            for _ in range(20):
                candidate = await page.query_selector(sel["send"])
                if candidate:
                    disabled = await candidate.get_attribute("disabled")
                    aria_disabled = await candidate.get_attribute("aria-disabled")
                    if disabled is None and aria_disabled != "true":
                        btn = candidate
                        break
                await asyncio.sleep(0.25)
            if btn:
                await self._safe_click(page, btn, f"{ai.upper()} 전송 버튼")
                self._post("system", f"✅ {ai.upper()}: 전송 버튼 클릭")
                sent = True
            else:
                self._post("system", f"⚠️ {ai.upper()}: 활성 전송 버튼 없음 — 키보드 전송 시도")
        except Exception:
            pass
        if not sent:
            await page.keyboard.press("Control+Enter")
            await asyncio.sleep(0.2)
            await page.keyboard.press("Enter")
            self._post("system", f"✅ {ai.upper()}: Ctrl+Enter/Enter 전송 시도")

        # 4) 새 응답이 끝날 때까지 대기하고 최종 텍스트 반환
        return await self._wait_until_done(page, sel, ai, before_response)

    async def _last_response_text(self, page: Page, sel: dict, ai: str = "") -> str:
        if ai == "perplexity":
            try:
                text = await page.evaluate(
                    """
                    () => {
                      const selectors = [
                        '[data-testid*="answer" i]',
                        'main [class*="prose"]',
                        'main [class*="markdown"]',
                        'main article',
                        'article'
                      ];
                      const candidates = [];
                      const seen = new Set();
                      for (const selector of selectors) {
                        for (const node of document.querySelectorAll(selector)) {
                          const rect = node.getBoundingClientRect();
                          const style = window.getComputedStyle(node);
                          if (!rect.width || !rect.height || style.display === 'none') continue;
                          const value = (node.innerText || node.textContent || '').trim();
                          if (value.length < 80 || seen.has(value)) continue;
                          seen.add(value);
                          candidates.push({ value, node });
                        }
                      }
                      candidates.sort((a, b) => a.value.length - b.value.length);
                      if (!candidates.length) return '';

                      const best = candidates[candidates.length - 1];
                      const scope = best.node.closest('article') || best.node.parentElement || best.node;
                      const links = [];
                      const seenUrls = new Set();
                      for (const anchor of scope.querySelectorAll('a[href]')) {
                        const href = anchor.href || '';
                        if (!href.startsWith('http') || seenUrls.has(href)) continue;
                        seenUrls.add(href);
                        const label = (anchor.innerText || anchor.textContent || '').trim();
                        links.push(`- ${label || '출처'}: ${href}`);
                      }
                      return links.length
                        ? `${best.value}\n\n출처 링크:\n${links.join('\n')}`
                        : best.value;
                    }
                    """
                )
                if text and text.strip():
                    return text.strip()
            except Exception:
                pass
        if ai == "gemini":
            try:
                text = await page.evaluate(
                    """
                    () => {
                      const selectors = [
                        'message-content',
                        'model-response',
                        '.model-response-text',
                        '[data-response-index]',
                        'div.markdown',
                        'div.prose'
                      ];
                      const seen = new Set();
                      const blocks = [];
                      for (const selector of selectors) {
                        for (const node of document.querySelectorAll(selector)) {
                          const text = (node.innerText || node.textContent || '').trim();
                          if (!text || seen.has(text)) continue;
                          seen.add(text);
                          blocks.push(text);
                        }
                      }
                      return blocks.length ? blocks[blocks.length - 1] : '';
                    }
                    """
                )
                if text and text.strip():
                    return text.strip()
            except Exception:
                pass
        for one_sel in [s.strip() for s in sel["response"].split(",") if s.strip()]:
            try:
                nodes = await page.query_selector_all(one_sel)
            except Exception:
                nodes = []
            if nodes:
                try:
                    text = (await nodes[-1].inner_text()).strip()
                except Exception:
                    text = ""
                if text:
                    return text
        return ""

    async def _fill_prompt(self, page: Page, input_el, prompt: str, ai: str) -> None:
        """긴 프롬프트를 줄 단위 입력 대신 Windows 클립보드 붙여넣기로 안정적으로 넣는다."""
        await page.keyboard.press("Control+A")
        await page.keyboard.press("Delete")
        await asyncio.sleep(0.15)

        if ai == "gemini":
            gemini_prompt = (
                "[빠르고 간결하게 답변해줘.]\n"
                "핵심 근거, 리스크, 실행 관점 위주로 정리해. "
                "단, 내용이 중간에 끊기지 않게 마지막 문장을 완결해.\n\n"
                f"{prompt}"
            )
            try:
                set_windows_clipboard_text(gemini_prompt)
                await page.keyboard.press("Control+V")
            except Exception:
                try:
                    await input_el.fill(gemini_prompt, timeout=60_000)
                except Exception:
                    await page.keyboard.insert_text(gemini_prompt)
            await asyncio.sleep(0.5)
            self._post("system", "✅ GEMINI: 프롬프트 입력 완료")
            return

        try:
            set_windows_clipboard_text(prompt)
            await page.keyboard.press("Control+V")
            await asyncio.sleep(0.8)
            if await self._input_contains_prompt(input_el, prompt):
                self._post("system", f"✅ {ai.upper()}: 클립보드 붙여넣기로 프롬프트 입력")
                return
            self._post(
                "system",
                f"ℹ️ {ai.upper()}: 붙여넣기 확인은 안 되지만 DOM 보강 없이 진행합니다. "
                "일부 웹 에디터는 붙여넣은 텍스트를 DOM에서 바로 읽을 수 없습니다."
            )
            return
        except Exception as e:
            self._post("system", f"⚠️ {ai.upper()}: 클립보드 붙여넣기 실패 — fill()로 재시도: {e}")

        try:
            await input_el.fill(prompt, timeout=60_000)
            await asyncio.sleep(0.5)
            await page.keyboard.press("End")
            await page.keyboard.insert_text(" ")
            await page.keyboard.press("Backspace")
            await asyncio.sleep(0.2)
            self._post("system", f"✅ {ai.upper()}: Playwright fill()로 프롬프트 입력")
            return
        except Exception as e:
            self._post("system", f"⚠️ {ai.upper()}: fill() 실패 — DOM 입력으로 재시도: {e}")

        await input_el.evaluate(
            """
            (el, text) => {
              el.focus();
              if ('value' in el) {
                el.value = text;
                el.dispatchEvent(new InputEvent('input', {
                  bubbles: true,
                  inputType: 'insertText',
                  data: text
                }));
                el.dispatchEvent(new Event('change', { bubbles: true }));
                return;
              }
              const selection = window.getSelection();
              const range = document.createRange();
              range.selectNodeContents(el);
              selection.removeAllRanges();
              selection.addRange(range);
              document.execCommand('delete', false, null);
              document.execCommand('insertText', false, text);
              el.dispatchEvent(new InputEvent('input', {
                bubbles: true,
                inputType: 'insertText',
                data: text
              }));
            }
            """,
            prompt,
        )
        await asyncio.sleep(0.5)
        if not await self._input_contains_prompt(input_el, prompt):
            self._post(
                "system",
                f"ℹ️ {ai.upper()}: 입력창 내용 검증을 건너뜁니다. "
                "일부 웹 에디터는 붙여넣은 텍스트를 DOM에서 바로 읽을 수 없습니다."
            )

    async def _input_contains_prompt(self, input_el, prompt: str) -> bool:
        try:
            text = await input_el.evaluate(
                """
                (el) => {
                  if ('value' in el) return el.value || '';
                  return el.innerText || el.textContent || '';
                }
                """
            )
        except Exception:
            return False
        normalized = (text or "").replace("\r\n", "\n").strip()
        expected = prompt.replace("\r\n", "\n").strip()
        if len(expected) < 80:
            return normalized == expected or expected in normalized
        return (
            len(normalized) >= int(len(expected) * 0.95)
            and expected[:80] in normalized
            and expected[-80:] in normalized
        )

    async def _safe_click(self, page: Page, element, label: str) -> None:
        """모달/온보딩 오버레이가 클릭을 가로막으면 닫고 한 번 더 시도."""
        try:
            await self._dismiss_blocking_overlays(page)
            await element.click(timeout=10_000)
            return
        except Exception as first_error:
            await self._dismiss_blocking_overlays(page, force=True)
            try:
                await element.click(timeout=10_000)
                return
            except Exception as second_error:
                raise RuntimeError(f"{label} 클릭 실패: {second_error}") from first_error

    async def _dismiss_blocking_overlays(self, page: Page, force: bool = False) -> None:
        """AI 서비스의 온보딩/팝업 모달을 best-effort로 닫는다."""
        # Escape로 닫히는 모달이 가장 흔하다. 부작용이 적어서 먼저 시도한다.
        try:
            await page.keyboard.press("Escape")
            await asyncio.sleep(0.25)
        except Exception:
            pass

        close_selectors = [
            'button[aria-label="Close"]',
            'button[aria-label="close"]',
            'button[aria-label="닫기"]',
            'button:has-text("닫기")',
            'button:has-text("Close")',
            'button:has-text("Not now")',
            'button:has-text("나중에")',
            'button:has-text("건너뛰기")',
            'button:has-text("Skip")',
            'button:has-text("Got it")',
            'button:has-text("확인")',
            'button:has-text("시작하기")',
            '[role="dialog"] button',
            '[data-state="open"] button',
        ]
        for selector in close_selectors:
            try:
                nodes = await page.query_selector_all(selector)
            except Exception:
                nodes = []
            for node in nodes[:4]:
                try:
                    if await node.is_visible():
                        await node.click(timeout=1_500)
                        await asyncio.sleep(0.35)
                        return
                except Exception:
                    continue

        if force:
            # 마지막 수단: 이벤트를 가로막는 전면 모달 컨테이너를 숨긴다.
            try:
                await page.evaluate(
                    """
                    () => {
                      const candidates = Array.from(document.querySelectorAll(
                        '[role="dialog"], [data-state="open"], .fixed.z-modal, .fixed.inset-0'
                      ));
                      for (const el of candidates) {
                        const style = window.getComputedStyle(el);
                        const rect = el.getBoundingClientRect();
                        const coversScreen = rect.width > window.innerWidth * 0.5 &&
                          rect.height > window.innerHeight * 0.5;
                        if ((style.position === 'fixed' || style.position === 'absolute') && coversScreen) {
                          el.style.pointerEvents = 'none';
                          el.style.display = 'none';
                        }
                      }
                    }
                    """
                )
                await asyncio.sleep(0.25)
            except Exception:
                pass

    def _pop_inject(self) -> str:
        """inject_queue에서 개입 메시지 꺼내기. 없으면 빈 문자열."""
        try:
            msg = self.inject_queue.get_nowait()
            self._post("system", f"💬 진행자 개입 주입: {msg}")
            return f"\n\n[진행자 개입 — 반드시 이 점을 고려해서 답변해줘]: {msg}\n"
        except Exception:
            return ""

    async def _try_google_signin(self, page: Page, ai: str) -> bool:
        """로그인 화면에서 Google 로그인을 열고, 설정된 계정이 있으면 선택한다."""

        # ── Step 1: "Continue with Google" 버튼 찾기 ──
        # Gemini는 accounts.google.com 로 직행, 별도 버튼 없음
        if ai != "gemini":
            google_btn_sels = [
                'button:has-text("Continue with Google")',
                'button:has-text("Google로 계속하기")',
                'button:has-text("Google 계정으로 계속")',
                'button:has-text("Google로 로그인")',
                'button[data-provider="google"]',
                'a:has-text("Continue with Google")',
            ]
            # ChatGPT는 "Log in" → "Continue with Google" 2단계
            entry_sels = ['button:has-text("Log in")', 'a:has-text("Log in")']

            async def _click_google_btn() -> bool:
                for s in google_btn_sels:
                    try:
                        el = await page.query_selector(s)
                        if el and await el.is_visible():
                            await el.click()
                            return True
                    except Exception:
                        continue
                return False

            if not await _click_google_btn():
                # "Log in" 버튼 먼저 눌러서 OAuth 화면 진입 시도
                for s in entry_sels:
                    try:
                        el = await page.query_selector(s)
                        if el and await el.is_visible():
                            await el.click()
                            await asyncio.sleep(1.5)
                            break
                    except Exception:
                        continue
                if not await _click_google_btn():
                    return False

            await asyncio.sleep(2.0)  # Google 팝업/리다이렉트 대기

        if not GOOGLE_ACCOUNT:
            return True

        # ── Step 2: Google 계정 선택 화면에서 설정된 계정 클릭 ──
        # Google AccountChooser 셀렉터 (다중 폴백)
        account_sels = [
            f'[data-email="{GOOGLE_ACCOUNT}"]',
            f'[data-identifier="{GOOGLE_ACCOUNT}"]',
            f'div:has-text("{GOOGLE_ACCOUNT}"):visible',
            f'li:has-text("{GOOGLE_ACCOUNT}")',
            f'[aria-label*="{GOOGLE_ACCOUNT}"]',
            f'a:has-text("{GOOGLE_ACCOUNT}")',
        ]
        # 새 팝업 탭이 열릴 수도 있어서 모든 페이지(탭) 검색
        context = page.context
        t0 = time.time()
        while time.time() - t0 < 15:
            for p in context.pages:
                if "accounts.google.com" in p.url or "google.com" in p.url:
                    for s in account_sels:
                        try:
                            el = await p.query_selector(s)
                            if el and await el.is_visible():
                                await el.click()
                                await asyncio.sleep(2.0)
                                return True
                        except Exception:
                            continue
            # 현재 페이지에서도 시도 (리다이렉트 케이스)
            for s in account_sels:
                try:
                    el = await page.query_selector(s)
                    if el and await el.is_visible():
                        await el.click()
                        await asyncio.sleep(2.0)
                        return True
                except Exception:
                    continue
            await asyncio.sleep(0.8)

        # 계정 선택 화면 못 찾아도 Google 버튼은 눌렀으니 True 반환 (사용자 수동 완료 기대)
        return True

    async def _try_download_chatgpt_attachment(
        self, timeout_sec: float = 90.0
    ) -> tuple[Optional[bytes], Optional[str]]:
        """ChatGPT Code Interpreter가 생성한 .docx 첨부 다운로드.
        (파일 바이트, 첨부 링크 URL)을 반환한다. 실패한 항목은 None."""
        page = self.pages["chatgpt"]
        await page.bring_to_front()

        candidates = [
            'a[href*="oaiusercontent.com"][download]',
            'a[href*="oaiusercontent.com"]',
            'a[download][href*=".docx"]',
            'a:has-text("AI_debate_report")',
            'a:has-text(".docx")',
            'button[aria-label*="Download" i]',
            '[data-testid*="download" i]',
        ]

        t0 = time.time()
        target = None
        while time.time() - t0 < 30 and target is None:
            for sel in candidates:
                try:
                    el = await page.query_selector(sel)
                    if el:
                        target = el
                        break
                except Exception:
                    pass
            if target is None:
                await asyncio.sleep(1.0)

        if target is None:
            return None, None

        link_url: Optional[str] = None
        try:
            link_url = await target.get_attribute("href")
        except Exception:
            pass

        try:
            async with page.expect_download(timeout=timeout_sec * 1000) as dl_info:
                await target.click()
            download = await dl_info.value
            download_path = await download.path()
            if not download_path:
                return None, link_url
            data = Path(download_path).read_bytes()
            return data, link_url
        except Exception:
            return None, link_url

    async def _try_download_claude_attachment(self, timeout_sec: float = 90.0) -> Optional[bytes]:
        """Claude 응답 영역에서 .docx 첨부 다운로드 버튼/링크를 찾아 클릭하고 파일 받기.
        Analysis 도구로 생성된 파일은 보통 아이콘+파일명+다운로드 버튼 UI로 나타남."""
        page = self.pages["claude"]
        await page.bring_to_front()

        # Claude UI에서 docx 첨부를 나타내는 후보 셀렉터들 (UI 변동 대비 다중)
        candidates = [
            'a[href$=".docx"]',
            'a[download][href*=".docx"]',
            'a[download][href*="report"]',
            'button[aria-label*="Download" i]:visible',
            'button[aria-label*="다운로드" i]',
            '[data-testid*="download" i]',
            'a:has-text("AI_debate_report")',
            'button:has-text("Download")',
            'button:has-text("다운로드")',
        ]

        # 응답 완료 후 첨부가 렌더링될 때까지 잠시 대기
        t0 = time.time()
        target = None
        while time.time() - t0 < 20 and target is None:
            for sel in candidates:
                try:
                    el = await page.query_selector(sel)
                except Exception:
                    el = None
                if el:
                    target = el
                    break
            if target is None:
                await asyncio.sleep(1.0)

        if target is None:
            return None

        # 다운로드 캡처
        try:
            async with page.expect_download(timeout=timeout_sec * 1000) as dl_info:
                await target.click()
            download = await dl_info.value
        except Exception:
            return None

        try:
            download_path = await download.path()
            if not download_path:
                return None
            data = Path(download_path).read_bytes()
            return data
        except Exception:
            return None

    async def _copy_last_gemini_response_text(self, page: Page, fallback_text: str) -> str:
        """Gemini 응답 하단의 복사 버튼을 눌러 최종 답변을 클립보드에서 읽는다."""
        sentinel = f"__AI_DEBATE_GEMINI_COPY_{time.time_ns()}__"
        try:
            set_windows_clipboard_text(sentinel)
        except Exception:
            sentinel = ""

        try:
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await asyncio.sleep(0.4)
            clicked_label = await page.evaluate(
                """
                () => {
                  const labelOf = (el) => [
                    el.innerText,
                    el.textContent,
                    el.getAttribute('aria-label'),
                    el.getAttribute('title'),
                    el.getAttribute('mattooltip'),
                    el.getAttribute('data-tooltip'),
                    el.getAttribute('data-testid')
                  ].filter(Boolean).join(' ').trim();

                  const candidates = [];
                  for (const el of document.querySelectorAll('button,[role="button"]')) {
                    const label = labelOf(el);
                    const lower = label.toLowerCase();
                    if (!lower.includes('copy') && !label.includes('복사') && !lower.includes('content_copy')) {
                      continue;
                    }
                    const style = window.getComputedStyle(el);
                    if (style.display === 'none' || style.visibility === 'hidden') {
                      continue;
                    }
                    const rect = el.getBoundingClientRect();
                    candidates.push({ el, label, top: rect.top, left: rect.left });
                  }

                  candidates.sort((a, b) => (a.top - b.top) || (a.left - b.left));
                  const target = candidates[candidates.length - 1];
                  if (!target) return '';
                  target.el.scrollIntoView({ block: 'center', inline: 'nearest' });
                  target.el.click();
                  return target.label || 'copy';
                }
                """
            )
        except Exception as e:
            self._post("system", f"ℹ️ GEMINI: 복사 버튼 클릭 실패 — DOM 텍스트 사용: {e}")
            return fallback_text

        if not clicked_label:
            self._post("system", "ℹ️ GEMINI: 복사 버튼을 찾지 못해 DOM 텍스트를 사용합니다.")
            return fallback_text

        await asyncio.sleep(0.8)
        try:
            copied = get_windows_clipboard_text().strip()
        except Exception:
            copied = ""

        if not copied or (sentinel and copied == sentinel):
            self._post("system", "ℹ️ GEMINI: 복사 후 클립보드가 바뀌지 않아 DOM 텍스트를 사용합니다.")
            return fallback_text

        fallback_norm = re.sub(r"\s+", " ", fallback_text or "").strip()
        copied_norm = re.sub(r"\s+", " ", copied).strip()
        if fallback_norm and len(copied_norm) < max(30, int(len(fallback_norm) * 0.5)):
            self._post("system", "ℹ️ GEMINI: 복사된 텍스트가 너무 짧아 DOM 텍스트를 사용합니다.")
            return fallback_text

        self._post("system", "✅ GEMINI: 하단 복사 버튼으로 응답 텍스트 확보")
        self._post("system", f"📋 GEMINI 복사본 ({len(copied)}자)\n\n{copied}")
        return copied

    async def _wait_until_done(
        self,
        page: Page,
        sel: dict,
        ai: str = "",
        before_text: str = "",
        hard_timeout: float = 420.0,
    ) -> str:
        """
        새 응답 텍스트가 시작되고 충분히 안정화될 때까지 대기한다.
        기존 마지막 답변(before_text)을 새 답변으로 오인하지 않도록 반드시 다른 텍스트를 기다린다.

        - Phase A: before_text와 다른 새 응답이 감지될 때까지 대기
        - Phase B: 텍스트가 오래 변하지 않고 Stop 버튼도 보이지 않을 때 완료
        """
        async def _stop_visible() -> bool:
            try:
                nodes = await page.query_selector_all(sel["stop"])
            except Exception:
                nodes = []
            for node in nodes:
                try:
                    if await node.is_visible():
                        return True
                except Exception:
                    continue
            return False

        started_at = time.time()
        first_change_at: Optional[float] = None
        last_change_at = time.time()
        prev = before_text.strip()
        cur = ""

        while time.time() - started_at < hard_timeout:
            await asyncio.sleep(1.0)
            cur = await self._last_response_text(page, sel, ai)

            if not cur:
                continue

            changed_from_before = cur.strip() != before_text.strip()
            if changed_from_before and first_change_at is None:
                first_change_at = time.time()
                self._post("system", "📝 새 응답 감지 — 완료될 때까지 대기 중…")

            if cur != prev:
                prev = cur
                last_change_at = time.time()
                continue

            if first_change_at is None:
                continue

            quiet_seconds = time.time() - last_change_at
            elapsed_after_start = time.time() - first_change_at
            if quiet_seconds >= 10 and elapsed_after_start >= 12:
                if await _stop_visible():
                    last_change_at = time.time()
                    continue
                # 한 번 더 확인해서 늦게 붙는 토큰을 방지한다.
                await asyncio.sleep(2.0)
                final_text = await self._last_response_text(page, sel, ai)
                if final_text == cur and not await _stop_visible():
                    if ai == "gemini":
                        return await self._copy_last_gemini_response_text(page, final_text)
                    return final_text
                prev = final_text
                last_change_at = time.time()

        raise TimeoutError("응답 대기 시간 초과 (7분)")


# ──────────────────────────── 싱글톤 ────────────────────────────
@st.cache_resource(show_spinner=False)
def get_worker() -> DebateWorker:
    w = DebateWorker()
    w.start()
    return w


HISTORY_DIR = Path(__file__).parent / "history"
START_TOPIC_PATH = Path(__file__).parent / "start_topic.json"


def save_history(topic: str, messages: list[dict], docx_filename: Optional[str]) -> str:
    """완료된 토론을 JSON으로 저장. 파일명 반환."""
    HISTORY_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    record = {
        "id": stamp,
        "topic": topic,
        "started_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "messages": messages,
        "docx_filename": docx_filename,
    }
    fn = f"{stamp}.json"
    (HISTORY_DIR / fn).write_text(
        json.dumps(record, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return fn


def list_history() -> list[dict]:
    """저장된 토론 메타데이터 목록 (최신순)."""
    if not HISTORY_DIR.exists():
        return []
    items = []
    for f in sorted(HISTORY_DIR.glob("*.json"), reverse=True):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            items.append({
                "id": data.get("id", f.stem),
                "topic": data.get("topic", "(제목 없음)"),
                "started_at": data.get("started_at", ""),
                "file": f.name,
                "docx_filename": data.get("docx_filename"),
            })
        except Exception:
            continue
    return items


def load_history(file_name: str) -> Optional[dict]:
    try:
        return json.loads((HISTORY_DIR / file_name).read_text(encoding="utf-8"))
    except Exception:
        return None


def normalize_debate_order(order: Optional[list[str]]) -> list[str]:
    return list(DEFAULT_DEBATE_ORDER)


def make_report_title(topic: str) -> str:
    compact = re.sub(r"\s+", " ", topic).strip()
    if not compact:
        return "AI 토론 최종 보고서"
    if len(compact) > 42:
        compact = compact[:42].rstrip() + "..."
    return f"{compact} - AI 토론 최종 보고서"


def make_report_filename(topic: str) -> str:
    compact = re.sub(r"\s+", "_", topic).strip("_")
    compact = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", compact)
    compact = compact[:48].strip("._ ")
    if not compact:
        compact = "AI_debate_report"
    return f"{compact}.docx"


def pop_start_topic_command() -> Optional[dict]:
    if not START_TOPIC_PATH.exists():
        return None
    try:
        payload = json.loads(START_TOPIC_PATH.read_text(encoding="utf-8"))
        topic = str(payload.get("topic", "")).strip()
        rounds = int(payload.get("rounds", 1))
        debate_order = normalize_debate_order(payload.get("debate_order"))
        no_local_records = bool(payload.get("no_local_records", True))
        if not topic:
            return None
        return {
            "topic": topic,
            "rounds": max(1, min(5, rounds)),
            "debate_order": debate_order,
            "no_local_records": no_local_records,
        }
    except Exception:
        return None


def clear_start_topic_command() -> None:
    try:
        START_TOPIC_PATH.unlink()
    except FileNotFoundError:
        pass
    except Exception:
        pass


@st.cache_resource(show_spinner=False)
def get_ngrok_url() -> Optional[str]:
    url: Optional[str] = None
    # start_tunnel.py가 남긴 사이드카 파일을 우선 사용
    sidecar = Path(__file__).parent / "ngrok_url.txt"
    if sidecar.exists():
        try:
            cached = sidecar.read_text(encoding="utf-8").strip()
            if cached:
                url = cached
        except Exception:
            pass

    # 사이드카가 없으면 직접 시도
    if not url and NGROK_AUTHTOKEN:
        try:
            ngrok.set_auth_token(NGROK_AUTHTOKEN)
            tunnel = ngrok.connect(STREAMLIT_PORT, "http")
            url = tunnel.public_url
        except Exception as e:
            try:
                print(f"[ngrok] failed: {e}".encode("ascii", "ignore").decode("ascii"))
            except Exception:
                pass

    return url


# ──────────────────────────── UI ────────────────────────────
ROLE_META = {
    "user":    ("🙋 USER",          "user"),
    "system":  ("⚙️ SYSTEM",        "assistant"),
    "perplexity": ("🔎 PERPLEXITY 팩트팩", "assistant"),
    "claude":  ("🟠 CLAUDE",        "assistant"),
    "chatgpt": ("🟢 CHATGPT",       "assistant"),
    "gemini":  ("🔵 GEMINI",        "assistant"),
    "summary": ("📝 통합 요약",     "assistant"),
}

DOCX_COLOR = {
    "user":    RGBColor(0xC2, 0x95, 0x00),
    "perplexity": RGBColor(0x20, 0x82, 0x8A),
    "claude":  RGBColor(0xD9, 0x77, 0x57),
    "chatgpt": RGBColor(0x10, 0xA3, 0x7F),
    "gemini":  RGBColor(0x42, 0x85, 0xF4),
    "summary": RGBColor(0x55, 0x55, 0x55),
}


def build_docx(messages: list[dict], topic_hint: str = "") -> bytes:
    """전체 토론 + 요약을 .docx 바이트로 변환."""
    doc = Document()
    title = doc.add_heading("🎭 다중 AI 토론 기록", 0)
    title.alignment = 1  # CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    if topic_hint:
        p = doc.add_paragraph()
        r = p.add_run(f"주제: {topic_hint}")
        r.bold = True
        r.font.size = Pt(13)

    doc.add_paragraph()  # spacer

    for m in messages:
        role = m["role"]
        if role == "system":
            continue  # 시스템 노이즈는 제외
        label, _ = ROLE_META.get(role, (role.upper(), "assistant"))
        h = doc.add_heading(level=2)
        run = h.add_run(label)
        if role in DOCX_COLOR:
            run.font.color.rgb = DOCX_COLOR[role]
        # 본문 (단락 단위 분할)
        for para in str(m.get("content", "")).split("\n"):
            if para.strip():
                doc.add_paragraph(para)
            else:
                doc.add_paragraph()

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _strip_md_inline(text: str) -> str:
    """Word 본문에 넣을 때 거슬리는 인라인 마크다운 기호를 정리."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)   # **bold**
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", text)  # *italic*
    text = re.sub(r"`([^`]+)`", r"\1", text)        # `code`
    return text


def build_docx_from_markdown(topic: str, md_text: str, report_title: Optional[str] = None) -> bytes:
    """Claude가 생성한 마크다운 문서를 Word(.docx)로 변환.
    헤더(#, ##, ###), 글머리표(-, *, •), 번호 목록(1.), 표(|...|) 지원."""
    doc = Document()
    h = doc.add_heading(report_title or make_report_title(topic), 0)
    h.alignment = 1

    sub = doc.add_paragraph()
    r = sub.add_run(f"주제: {topic}")
    r.bold = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.add_run(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    doc.add_paragraph()

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.lstrip()

        # 표 (간단한 markdown 파이프 테이블)
        if stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]:
            # 다음 라인이 구분선이면 진짜 표
            sep = lines[i + 1].strip() if i + 1 < len(lines) else ""
            if re.match(r"^\|[\s\-:|]+\|$", sep):
                # 표 수집
                rows: list[list[str]] = []
                rows.append([c.strip() for c in stripped.strip("|").split("|")])
                i += 2  # header + sep
                while i < len(lines):
                    row = lines[i].strip()
                    if row.startswith("|") and row.endswith("|"):
                        rows.append([c.strip() for c in row.strip("|").split("|")])
                        i += 1
                    else:
                        break
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ""
                        table.rows[ri].cells[ci].text = _strip_md_inline(cell_text)
                doc.add_paragraph()
                continue

        if not stripped:
            doc.add_paragraph()
        elif stripped.startswith("### "):
            doc.add_heading(_strip_md_inline(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_strip_md_inline(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_strip_md_inline(stripped[2:]), level=1)
        elif stripped.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(_strip_md_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(_strip_md_inline(re.sub(r"^\d+\.\s", "", stripped)),
                              style="List Number")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(_strip_md_inline(stripped[2:]))
            p.paragraph_format.left_indent = Pt(18)
        elif stripped.startswith("---"):
            doc.add_paragraph()
        else:
            doc.add_paragraph(_strip_md_inline(stripped))
        i += 1

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def send_completion_notifications(
    topic: str,
    filename: Optional[str],
    gpt_chat_url: Optional[str] = None,
    persisted: bool = True,
) -> list[str]:
    """토론 완료 사실만 알림. 결과 본문과 파일은 외부 알림에 싣지 않는다."""
    sent: list[str] = []
    message = build_completion_notice(topic, filename, gpt_chat_url, persisted)
    if TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID:
        send_telegram_notice(message)
        sent.append("Telegram")
    return sent


def build_completion_notice(
    topic: str,
    filename: Optional[str],
    gpt_chat_url: Optional[str] = None,
    persisted: bool = True,
) -> str:
    file_line = ""
    if filename:
        file_line = (
            f"\n저장 파일: out/{filename}"
            if persisted
            else f"\n다운로드 파일명: {filename} (PC에 자동 저장하지 않음)"
        )
    chat_line = f"\nChatGPT 대화: {gpt_chat_url}" if gpt_chat_url else ""
    return (
        "AI 토론이 완료되었습니다.\n"
        "결과 내용은 이 알림에 포함하지 않았습니다.\n"
        "PC의 Streamlit 화면 또는 각 AI 채팅 탭에서 직접 확인하세요.\n\n"
        f"주제: {topic}"
        f"{file_line}"
        f"{chat_line}"
    )


def send_telegram_notice(message: str) -> None:
    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    data = json.dumps(
        {
            "chat_id": TELEGRAM_CHAT_ID,
            "text": message,
            "disable_web_page_preview": True,
        },
        ensure_ascii=False,
    ).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={"Content-Type": "application/json; charset=utf-8"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=20) as res:
        if res.status >= 400:
            raise RuntimeError(f"Telegram HTTP {res.status}")


CUSTOM_CSS = """
<style>
/* ── ChatGPT풍 라이트 테마 ── */
.stApp { background: #FFFFFF; }
.main .block-container {
    max-width: 768px;
    padding-top: 1rem;
    padding-bottom: 7rem;  /* 하단 고정 입력창 공간 */
}

/* Sidebar: ChatGPT처럼 옅은 회색, 가는 경계선 */
[data-testid="stSidebar"] {
    background: #F9F9F9;
    border-right: 1px solid #ECECEC;
}
[data-testid="stSidebar"] .stMarkdown h3 { font-size: 1.05rem; }

/* 채팅 메시지: 카드 느낌 제거, ChatGPT처럼 본문 흐름형 */
[data-testid="stChatMessage"] {
    background: transparent;
    border: none;
    box-shadow: none;
    padding: 4px 0;
    margin-bottom: 2px;
}

/* 사용자 메시지만 말풍선 (ChatGPT의 회색 캡슐) */
[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) {
    background: #F4F4F4;
    border-radius: 18px;
    padding: 10px 16px;
}

/* 하단 채팅 입력창: 둥근 캡슐 + 은은한 그림자 */
[data-testid="stChatInput"] {
    border: 1px solid #D9D9E3;
    border-radius: 26px;
    box-shadow: 0 2px 12px rgba(0,0,0,0.06);
    background: #FFFFFF;
}
[data-testid="stChatInput"] textarea {
    background: transparent !important;
}

/* 빈 화면 히어로 (ChatGPT 첫 화면 느낌) */
.hero {
    text-align: center;
    margin-top: 18vh;
    color: #0D0D0D;
}
.hero h2 { font-weight: 600; letter-spacing: -0.3px; }
.hero p { color: #8E8EA0; font-size: 0.95rem; }

/* 최종 보고서 카드 */
.report-card {
    border: 1px solid #ECECEC;
    border-radius: 16px;
    padding: 16px 20px;
    background: #FAFAFA;
    margin: 10px 0 16px 0;
}

/* Buttons */
.stButton > button, .stDownloadButton > button {
    border-radius: 12px;
    font-weight: 600;
}

.stCaption, [data-testid="stCaptionContainer"] { color: #8E8EA0 !important; }

code {
    background: #F4F4F4 !important;
    color: #1F2328 !important;
    padding: 2px 6px;
    border-radius: 4px;
}
</style>
"""

ROLE_AVATARS = {
    "user": "🙋",
    "system": "⚙️",
    "perplexity": "🔎",
    "claude": "🟠",
    "chatgpt": "🟢",
    "gemini": "🔵",
    "summary": "📝",
}


def main() -> None:
    st.set_page_config(
        page_title="AI 토론",
        page_icon="🎭",
        layout="centered",
        initial_sidebar_state="expanded",
    )
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

    ngrok_url = get_ngrok_url()
    worker = get_worker()
    status, err, messages, last_docx = worker.snapshot()

    start_cmd = pop_start_topic_command()
    if start_cmd and status == "ready":
        worker.submit_topic(
            start_cmd["topic"],
            start_cmd["rounds"],
            start_cmd.get("debate_order"),
            no_local_records=bool(start_cmd.get("no_local_records", True)),
        )
        clear_start_topic_command()
        st.rerun()

    with st.sidebar:
        st.markdown("### 🎭 AI 토론")
        badge = {"booting": "🟡 부팅 중", "ready": "🟢 대기", "running": "🔵 토론 진행 중", "error": "🔴 오류"}
        st.caption(badge.get(status, status))
        if err:
            st.error(err)
        if ngrok_url:
            st.caption(f"🌐 외부 접속: `{ngrok_url}`")
        else:
            st.caption("🌐 ngrok 비활성 — `.env`에 `NGROK_AUTHTOKEN` 설정 시 외부 접속이 열립니다.")

        # ── 결과 내보내기 ──
        if last_docx:
            st.divider()
            st.markdown("**📦 결과 내보내기**")
            st.caption("완료 알림은 Telegram으로만 발송됩니다. 결과 내용과 파일은 알림에 포함하지 않습니다.")
            st.download_button(
                "⬇️ 최종 Word(.docx) 다운로드",
                data=last_docx["bytes"],
                file_name=last_docx["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            # 토론 전체 백업본도 별도 제공
            try:
                transcript_bytes = build_docx(messages, topic_hint=last_docx.get("topic", ""))
                st.download_button(
                    "🗂 토론 전체 기록 다운로드 (백업)",
                    data=transcript_bytes,
                    file_name=f"AI_debate_full_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception:
                pass

        # ── 히스토리 ──
        st.divider()
        st.markdown("**🕒 토론 히스토리**")
        history_items = list_history()
        if not history_items:
            st.caption("(아직 저장된 토론이 없어요)")
            selected_file = None
        else:
            labels = ["▶ 현재 세션"] + [
                f"{h['started_at']} — {h['topic'][:34]}{'…' if len(h['topic']) > 34 else ''}"
                for h in history_items[:50]
            ]
            sel_idx = st.selectbox(
                "지난 토론 보기",
                options=range(len(labels)),
                format_func=lambda i: labels[i],
                index=0,
                key="history_idx",
            )
            selected_file = None if sel_idx == 0 else history_items[sel_idx - 1]["file"]

        st.divider()
        st.markdown(
            "**📌 첫 실행 시 체크리스트**\n"
            "1. `.env`에 ngrok 토큰 입력\n"
            "2. 새 크롬 창에서 Perplexity / Gemini / ChatGPT / Claude 로그인 및 보안 확인\n"
            "3. 캡챠 뜨면 크롬 창에서 직접 풀기"
        )

    # ── 메인 영역: 히스토리 선택됐으면 그 메시지 표시, 아니면 현재 세션 ──
    if selected_file:
        record = load_history(selected_file)
        if record:
            st.info(f"📜 **과거 토론 보기**: {record['topic']}  ·  {record['started_at']}")
            view_messages = record["messages"]
            # 과거 토론의 Word 파일이 out/에 남아있으면 다운로드도 제공
            if record.get("docx_filename"):
                docx_path = Path(__file__).parent / "out" / record["docx_filename"]
                if docx_path.exists():
                    st.download_button(
                        f"⬇️ {record['docx_filename']} 다운로드",
                        data=docx_path.read_bytes(),
                        file_name=record["docx_filename"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
        else:
            st.warning("히스토리 파일을 읽지 못했습니다.")
            view_messages = []
    else:
        if status == "ready" and last_docx:
            if last_docx.get("persisted"):
                st.info("이전 토론은 히스토리와 결과 다운로드 영역에 보관했습니다. 새 주제는 아래 입력창에서 바로 시작하세요.")
            else:
                st.info("완료된 토론은 현재 메모리에만 있습니다. 서버를 재시작하면 사라집니다.")
            view_messages = []
        else:
            view_messages = messages

    if not view_messages:
        st.markdown(
            '<div class="hero">'
            "<h2>무엇이든 토론에 부쳐 보세요</h2>"
            "<p>Perplexity가 근거를 수집한 뒤 Gemini · ChatGPT · Claude가 토론을 시작합니다.<br>"
            "예) 한국 부동산의 향후 2년 전망은 상승일까 하락일까?</p>"
            "</div>",
            unsafe_allow_html=True,
        )
    else:
        for m in view_messages:
            role = m["role"]
            label, bubble_role = ROLE_META.get(role, (role.upper(), "assistant"))
            with st.chat_message(bubble_role, avatar=ROLE_AVATARS.get(role)):
                if role != "user":
                    st.markdown(f"**{label}**")
                st.markdown(m["content"])

    # ── 최종 보고서 카드: 다운로드 + ChatGPT 문서 링크 ──
    if last_docx and not selected_file:
        with st.container(border=True):
            st.markdown(f"#### 📄 {last_docx.get('title') or '최종 보고서'}")
            if last_docx.get("gpt_chat_url"):
                st.markdown(f"💬 [문서를 만든 ChatGPT 대화 열기]({last_docx['gpt_chat_url']})")
            if last_docx.get("gpt_link"):
                st.markdown(f"🔗 [ChatGPT 생성 문서 링크]({last_docx['gpt_link']})")
                st.caption("ChatGPT 파일 링크는 시간이 지나면 만료될 수 있어요. 다운로드 버튼이 안전합니다.")
            st.download_button(
                "⬇️ 최종 Word(.docx) 다운로드",
                data=last_docx["bytes"],
                file_name=last_docx["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="report_card_docx",
            )

    # ── 하단 토론 설정: 새 토론 시작 직전에 조율 ──
    with st.container(border=True):
        st.markdown("**토론 설정**")
        no_local_records = st.checkbox(
            "이 앱의 질문·답변 기록을 PC에 남기지 않기",
            value=True,
            key="no_local_records",
            disabled=status != "ready",
            help=(
                "체크하면 history, logs, out 폴더에 이번 토론의 질문·답변과 Word 파일을 "
                "저장하지 않습니다. 화면 다운로드는 가능하며 서버를 재시작하면 메모리 결과는 사라집니다."
            ),
        )
        if no_local_records:
            st.caption(
                "로컬 비공개 모드: 이 앱의 history/logs/out 저장만 차단합니다. "
                "자동화 Chrome 프로필의 방문 기록·캐시와 각 AI 서비스 서버의 대화 기록은 "
                "각 브라우저·서비스 설정에서 별도로 관리해야 합니다."
            )
        c1, c2 = st.columns([1, 2])
        with c1:
            rounds = st.slider(
                "토론 세트 수",
                min_value=1,
                max_value=5,
                value=int(st.session_state.get("rounds_control", 1)),
                key="rounds_control",
                disabled=status != "ready",
                help="선택한 순서로 몇 번 왕복할지 정합니다. 마지막 AI가 최종 보고서를 작성합니다.",
            )
        with c2:
            st.text_input(
                "토론 순서",
                value=(
                    "Perplexity 근거 수집 → Gemini → ChatGPT → Claude "
                    "→ GPT 최종 반대자 검토 → Claude 수정 최종안"
                ),
                disabled=True,
                help=(
                    "Perplexity는 사실과 출처를 수집하고, Gemini는 근거 기반으로 발산하며, "
                    "GPT는 검증하고 Claude는 수렴합니다. Claude 초안 뒤 GPT가 최종 반대자 "
                    "검토를 하고 Claude가 수정 최종안을 작성합니다."
                ),
            )
        debate_order = list(DEFAULT_DEBATE_ORDER)
        if status == "running":
            st.caption("토론 진행 중에는 아래 채팅창에 쓰는 메시지가 다음 발화 전에 Human-in-the-loop 개입으로 들어갑니다.")

    # ── 하단 채팅 입력 (ChatGPT 스타일): 대기 중엔 주제, 진행 중엔 개입 ──
    if status == "ready":
        user_msg = st.chat_input("토론 주제를 입력하세요…")
        if user_msg and user_msg.strip():
            worker.submit_topic(
                user_msg.strip(),
                rounds,
                debate_order,
                no_local_records=no_local_records,
            )
            st.rerun()
    elif status == "running":
        user_msg = st.chat_input("토론에 개입할 메시지 — 다음 발화 전에 반영됩니다…")
        if user_msg and user_msg.strip():
            worker.inject_message(user_msg.strip())
            st.rerun()
    else:
        st.chat_input("준비 중입니다…", disabled=True)

    # 진행 중이면 자동 새로고침
    if status in ("booting", "running") and not selected_file:
        time.sleep(1.5)
        st.rerun()


if __name__ == "__main__":
    main()
